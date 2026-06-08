from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "深圳市养老护理院景观改造建筑策划报告_A4详细版.docx"
SOURCE_PPT = r"C:\Users\19727\Desktop\建筑策划\建筑策划汇报.pptx"


def set_run(run, size=10.5, bold=False, color="1F2937", name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_para_spacing(paragraph, after=6, before=0, line=1.16):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def write_cell(cell, text, size=8.0, bold=False, color="1F2937", align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    set_para_spacing(p, after=0, line=1.05)
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text))
    set_run(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell)


def add_table(doc, headers, rows, widths, header_fill="E8EEF5", size=7.9, align_first=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        shade_cell(cell, header_fill)
        write_cell(cell, header, size=size, bold=True, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.width = Inches(widths[idx])
            align = WD_ALIGN_PARAGRAPH.CENTER if align_first and idx == 0 else None
            write_cell(cell, value, size=size, align=align)
    doc.add_paragraph()
    return table


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    set_para_spacing(p, before=14, after=8, line=1.12)
    for r in p.runs:
        set_run(r, size=16, bold=True, color="1F4D78")


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    set_para_spacing(p, before=10, after=5, line=1.12)
    for r in p.runs:
        set_run(r, size=13, bold=True, color="2E74B5")


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    set_para_spacing(p, before=8, after=4, line=1.12)
    for r in p.runs:
        set_run(r, size=11.5, bold=True, color="0B2545")


def para(doc, text, lead=None):
    p = doc.add_paragraph()
    set_para_spacing(p, after=6, line=1.18)
    if lead and text.startswith(lead):
        r = p.add_run(lead)
        set_run(r, bold=True, color="0B2545")
        r = p.add_run(text[len(lead):])
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para_spacing(p, after=4, line=1.14)
        r = p.add_run(item)
        set_run(r, size=10.2)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_para_spacing(p, after=4, line=1.14)
        r = p.add_run(item)
        set_run(r, size=10.2)


def add_hyperlink(paragraph, text, url):
    rel_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_note_table(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.rows[0].cells[0].width = Inches(6.45)
    shade_cell(table.rows[0].cells[0], "F4F6F9")
    cell_margins(table.rows[0].cells[0], top=120, bottom=120, start=140, end=140)
    p = table.rows[0].cells[0].paragraphs[0]
    set_para_spacing(p, after=2, line=1.14)
    r = p.add_run(title)
    set_run(r, size=10.2, bold=True, color="0B2545")
    r = p.add_run(body)
    set_run(r, size=9.6, color="374151")
    doc.add_paragraph()


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.16
    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        styles[style_name].font.name = "Microsoft YaHei"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("深圳市养老护理院景观改造建筑策划报告 | A4详细版")
    set_run(r, size=8, color="6B7280")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("基于积极老龄化理念的疗愈型适老化景观更新研究")
    set_run(r, size=8, color="6B7280")
    return doc


def add_references(doc):
    h1(doc, "参考文献与资料来源")
    refs = [
        ("深圳市南山区人民政府：深圳市养老护理院机构信息页。", "https://www.szns.gov.cn/bsfw/nsqyljgmc/szsylhly/content/post_9447864.html"),
        ("深圳市人民政府：深圳市养老服务设施布局专项规划（2025-2035年）公开信息。", "https://www.sz.gov.cn/cn/xxgk/zfxxgj/ghjh/csgh/zxgh/content/post_12470457.html"),
        ("World Health Organization. Active Ageing: A Policy Framework. 2002.", "https://extranet.who.int/agefriendlyworld/wp-content/uploads/2014/06/WHO-Active-Ageing-Framework.pdf"),
        ("World Health Organization. Global Age-friendly Cities: A Guide. 2007.", "https://iris.who.int/handle/10665/43755"),
        ("Whole Building Design Guide. Architectural Programming.", "https://www.wbdg.org/design-disciplines/architectural-programming"),
        ("Ulrich, R. S. View through a window may influence recovery from surgery. Science, 1984.", "https://doi.org/10.1126/science.6143402"),
        ("Marcus, C. C., & Sachs, N. A. Therapeutic Landscapes: An Evidence-Based Approach to Designing Healing Gardens and Restorative Outdoor Spaces. Wiley, 2014.", "https://www.wiley.com/en-us/Therapeutic+Landscapes%3A+An+Evidence+Based+Approach+to+Designing+Healing+Gardens+and+Restorative+Outdoor+Spaces-p-9781118421109"),
        ("住房和城乡建设部：建筑与市政工程无障碍通用规范 GB 55019-2021。", "https://www.mohurd.gov.cn/"),
        ("住房和城乡建设部：老年人照料设施建筑设计标准 JGJ 450-2018。", "https://www.mohurd.gov.cn/"),
        ("住房和城乡建设部：城市居住区规划设计标准 GB 50180-2018。", "https://www.mohurd.gov.cn/gongkai/zhengce/zhengcefilelib/201811/20181130_238590.html"),
        (f"源PPT：建筑策划汇报.pptx，2026-05-25，深圳市养老护理院景观改造建筑策划报告。{SOURCE_PPT}", ""),
    ]
    for idx, (label, url) in enumerate(refs, 1):
        p = doc.add_paragraph()
        set_para_spacing(p, after=4, line=1.08)
        r = p.add_run(f"[R{idx}] {label} ")
        set_run(r, size=8.6)
        if url:
            add_hyperlink(p, url, url)


def build():
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(title, after=3)
    r = title.add_run("深圳市养老护理院景观改造建筑策划报告")
    set_run(r, size=22, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(subtitle, after=14)
    r = subtitle.add_run("基于积极老龄化理念的疗愈型适老化景观更新研究")
    set_run(r, size=12.5, color="2E74B5")

    add_table(
        doc,
        ["项目", "说明"],
        [
            ("报告定位", "建筑策划/景观更新策划报告。用于方案设计前明确价值排序、空间任务、运营机制和技术控制要求。"),
            ("研究对象", "深圳市养老护理院院内庭院、街心公园、室外慢行系统、集中活动区域及院内外衔接界面。"),
            ("核心目标", "从“被动照护”转向“积极生活”，以安全、疗愈、社交与认知友好重构养老院户外日常。"),
            ("资料依据", f"源PPT：{SOURCE_PPT}；用户补充参考图；公开政策、标准与循证研究文献。"),
            ("面积口径", "本阶段未取得CAD测绘文件，面积采用总更新面积 A 的比例控制，施工图阶段复核为平方米。"),
        ],
        [1.25, 5.15],
        size=8.8,
    )
    add_note_table(
        doc,
        "策划结论：",
        "本项目不应被理解为庭院美化，而应被定义为“养老护理院户外康养生活系统更新”。策划重点是把老人真实行为、护理工作流、社区共享和长期维护机制同时纳入空间设计任务。"
    )
    doc.add_page_break()

    h1(doc, "1. 选址")
    h2(doc, "1.1 区位与项目属性")
    para(doc, "项目位于深圳市南山区桃源街道龙珠七路与龙苑路交界处的深圳市养老护理院及其相邻街心公园界面。公开资料显示，深圳市养老护理院用地面积约1万平方米，建筑面积约4万平方米，床位约800张，是承担长期照护、护理康复和医养结合服务的城市养老机构。")
    para(doc, "本次策划属于既有养老设施的存量更新，不涉及新建选址论证中的大尺度用地比选，而是针对“既有院区与邻近公共绿地如何转化为康养支持系统”展开。其价值在于：在高密度城区中，通过较小尺度的景观和设施更新，提高老人外出频率、活动安全性、家属探访质量和机构公共形象。")
    h2(doc, "1.2 选址价值判断")
    add_table(
        doc,
        ["维度", "现状判断", "策划意义", "控制要求"],
        [
            ("可达性", "院区位于城市居住片区，具备家属探访与社区资源导入基础。", "可将机构从封闭照护单元转为健康社区节点。", "入口识别、无障碍衔接、院内外流线分级。"),
            ("安全性", "院内相对可控，街心公园存在高差、消防口打断与线性空间不连续。", "安全是老人外出、康复和社交的前提。", "连续扶手、防滑、坡道、夜间照明、可视看护。"),
            ("疗愈性", "庭院、树荫、街心绿地已有自然资源，但体验以观看为主。", "可转化为情绪安抚、认知刺激和康复训练媒介。", "五感植物、声景、触觉材料、园艺疗法。"),
            ("运营性", "老人已有散步、晒太阳、交谈、锻炼、探访等高频行为。", "行为基础明确，适合形成稳定课程和维护机制。", "活动日历、护理记录、志愿机制、低维护设施。"),
        ],
        [0.85, 2.05, 1.85, 1.65],
        size=7.8,
    )
    h2(doc, "1.3 选址边界与更新范围")
    bullets(doc, [
        "院内庭院：以老人日常使用为主，强调护理可控、安全闭环、安静疗愈和适老家具。",
        "街心公园：以外出活动、社区共享和认知友好宣教为主，强调可管理开放和城市界面改善。",
        "室外慢行系统：作为串联院内庭院、入口、集中活动区和街心公园的康养骨架。",
        "集中活动区域：作为群体活动、家属探访、志愿服务和节庆展示的复合节点。",
    ])

    h1(doc, "2. 项目背景研究")
    h2(doc, "2.1 养老设施更新趋势")
    para(doc, "随着养老服务从机构床位供给转向全生命周期健康支持，养老院户外空间的角色正在发生变化。传统护理院外部环境常被视为绿化配套，但积极老龄化和年龄友好理念要求空间支持老人健康、参与、尊严和社会连接。疗愈景观研究进一步提示，自然环境、可参与活动、可控制的感官刺激和社会支持，对情绪稳定、康复动力和日常生活质量均有积极意义。")
    h2(doc, "2.2 项目问题背景")
    add_table(
        doc,
        ["问题背景", "具体表现", "深层原因", "策划回应"],
        [
            ("照护模式转型", "老人不只需要被照护，也需要活动、社交和被看见。", "养老机构空间仍偏向管理和通行，生活场景不足。", "构建参与式户外生活网络。"),
            ("身体机能差异", "自理、介助、介护老人对坡度、休息、扶手、看护的需求不同。", "空间未充分分级，设施未形成连续系统。", "安全闭环、分级开放、护理可视。"),
            ("认知与情绪需求", "方向感下降、焦虑、孤独和活动意愿下降影响户外使用。", "导视、节点记忆、感官刺激和交往场景不足。", "认知友好导视、记忆疗愈、多感官花园。"),
            ("运营维护压力", "一次性改造容易在后期变成维护负担。", "缺少模块化、低维护和活动机制设计。", "针灸式微更新、低维护植物、活动日历。"),
        ],
        [1.0, 1.8, 1.8, 1.8],
        size=7.7,
    )
    h2(doc, "2.3 策划目标")
    add_table(
        doc,
        ["目标", "内涵", "设计转译", "运营转译"],
        [
            ("安全性", "让老人走得稳、找得到路、走累能停、需要帮助能被看见。", "连续慢行、扶手、防滑、照明、紧急求助、休息点。", "风险巡检、跌倒记录复盘、护理陪同路线。"),
            ("疗愈性", "通过自然、园艺、声景、芳香和康复活动促进身心状态。", "五感花园、冥想节点、园艺床、芳香植物、地滚球。", "园艺课、康复小组、情绪观察记录。"),
            ("社交性", "让老人、家属、护理人员、志愿者和社区形成温和连接。", "半围合座椅、探访节点、共聚之园、展示界面。", "生日会、茶话会、节庆展示、志愿陪伴。"),
            ("认知友好", "降低迷失和焦虑，强化路径记忆和节点识别。", "色彩分区、图形导视、老物件展示、认知体感墙。", "认知症友好宣教、记忆叙事活动。"),
        ],
        [0.85, 1.95, 1.9, 1.7],
        size=7.7,
    )

    h1(doc, "3. 收集与分析信息")
    h2(doc, "3.1 信息收集框架")
    para(doc, "建筑策划的信息收集不应只停留在现场照片和功能清单，而应同时回答四个问题：谁在使用、何时使用、为什么使用、空间如何支持或阻碍行为。基于源PPT和会议参考图，本报告将信息分为场地、行为、人群、运营、政策五类。")
    add_table(
        doc,
        ["信息类型", "已掌握内容", "分析目的", "后续深化资料"],
        [
            ("场地现状", "院内庭院、街心公园、慢行系统、集中活动区照片与问题归纳。", "识别安全、体验、认知、运营缺口。", "CAD、竖向标高、坡度、消防口、排水、照明点位。"),
            ("老人行为", "散步、晒太阳、交谈、锻炼、探访等高频行为及时间规律。", "形成行为-空间-设施-运营转译。", "分时段观察、停留热力、跌倒风险、护理陪同记录。"),
            ("使用人群", "自理、介助、介护老人，家属，护理人员，康复师，社工，志愿者。", "确定活动强度、开放层级和看护距离。", "行动能力分级、认知症比例、护理工作流。"),
            ("运营条件", "现状有活动需求，但课程场景、维护接口和社区机制不足。", "判断哪些空间能长期被使用和维护。", "活动日历、维护预算、志愿者协议、设备巡检制度。"),
            ("政策标准", "积极老龄化、年龄友好、疗愈景观、无障碍和养老设施标准。", "为设计指标和任务书提供底线。", "地方审查意见、消防审查、院方管理规定。"),
        ],
        [0.95, 1.9, 1.75, 1.8],
        size=7.4,
    )
    h2(doc, "3.2 现状问题诊断")
    add_table(
        doc,
        ["问题", "表层表现", "对老人行为的影响", "策划方向"],
        [
            ("安全缺口", "扶手、坡道、座椅借力点、夜间照明未形成连续系统。", "降低老人独立外出意愿，增加护理陪同压力。", "补齐安全设施闭环，优先处理高差、湿滑和夜间风险。"),
            ("体验缺口", "景观以观看为主，触摸、声音、园艺和运动参与不足。", "停留时间短，活动重复性弱，疗愈效益有限。", "导入多感官花园、园艺疗法和低强度运动。"),
            ("认知缺口", "导视、图形识别、节点记忆和色彩引导不足。", "方向感下降老人易焦虑，路径选择负担高。", "以图形、色彩、主题节点和记忆展陈降低复杂度。"),
            ("社交缺口", "半围合、小尺度交流和家属探访场景不足。", "老人交往依赖偶遇，家属停留质量不稳定。", "设置面对面座椅、探访花园和共聚活动节点。"),
            ("运营缺口", "缺少固定课程、维护制度和可替换展陈。", "改造后可能回到“只看不参与”。", "把空间节点与护理、康复、社工、志愿活动绑定。"),
        ],
        [0.9, 1.8, 1.85, 1.85],
        size=7.5,
    )
    h2(doc, "3.3 行为需求分析")
    add_table(
        doc,
        ["行为", "时间/情境", "空间需求", "设施配置", "运营触发"],
        [
            ("散步", "早餐后、午饭后、傍晚。", "连续、可识别、可中止。", "环线、扶手、休息点、遮阴。", "康复步行记录、晨间散步。"),
            ("晒太阳", "午前后、冬季南向停留。", "日照、避风、可移动座椅。", "半围合座椅、暖色铺装。", "陪伴聊天、节气活动。"),
            ("交谈", "庭院、树下、家属探访。", "面对面、小尺度、护理可视。", "围合座椅、茶话节点。", "生日会、志愿陪伴。"),
            ("锻炼", "自理老人主动，介助老人陪同。", "低冲击、低风险、可看护。", "地滚球、平衡训练、环形跑道。", "康复小组、运动打卡。"),
            ("认知刺激", "散步过程、探访过程。", "低复杂度、可回忆、可触摸。", "认知墙、老物件、色彩导视。", "记忆叙事、认知友好宣教。"),
            ("园艺参与", "课程、采摘、志愿活动。", "可坐可站、安全工具。", "抬高种植床、可食植物。", "园艺课、植物认养。"),
        ],
        [0.75, 1.25, 1.45, 1.45, 1.55],
        size=7.1,
    )

    h1(doc, "4. 工作会议")
    h2(doc, "4.1 会议组织方法")
    para(doc, "工作会议的作用是把模糊的愿望转化为可设计、可审查、可运营的任务。参考用户提供的会议记录样式，本报告将会议内容反推为“价值议题-目标发言-事实信息-需求发言-会议归纳”，再转译为成果矩阵。")
    numbered(doc, [
        "启动会：明确范围、更新边界、资料清单、施工影响和目标人群。",
        "现状诊断会：核对高频路径、跌倒风险点、日照停留点、护理盲区和家属探访需求。",
        "概念策划会：确认四位一体目标、五园一廊一径、街心公园南北分区和运营场景。",
        "技术协调会：复核坡道、扶手、铺装、照明、排水、消防口、植物维护和设施安全。",
        "运营落地会：确定园艺、康复、认知友好宣教、志愿服务和社区共享机制。",
    ])
    h2(doc, "4.2 会议记录反推")
    add_table(
        doc,
        ["价值议题", "目标类发言（反推）", "事实信息类发言（反推）", "需求类发言（反推）"],
        [
            ("环境", "“希望增加聚集地，使老人愿意交流”；“晒太阳的位置要舒服，也要有树荫”。", "植物层次不足；导视不清；中庭与南向空间停留多；街心公园活力不足。", "绿荫、日照、芳香植物、声景、休息座椅、低维护植物。"),
            ("功能", "“空间要服务康复和日常照护”；“家属探访时要有可坐、可看的地方”。", "散步、交谈、锻炼、探访高频；集中活动区座椅和日照不足。", "地滚球、园艺课程、探访座椅、活动展示、护理陪同点。"),
            ("交通", "“老人要知道往哪走，走累了能停”；“轮椅和陪护通行不能被台阶卡住”。", "扶手不连续；台阶和高差阻隔；消防口打断街心公园线性空间。", "无障碍坡道、连续扶手、防滑铺装、约50米休息点、夜间照明。"),
            ("经济", "“施工影响要小，后期维护不能太复杂”。", "既有机构运营不中断；护理院施工扰动敏感；维护人手有限。", "分期实施、低成本微更新、模块化设施、可替换展陈。"),
            ("文化", "“老人需要被看见，也需要和社区有联系”。", "机构封闭感较强；记忆展示、认知症友好宣教和社区参与不足。", "老物件、老照片、节庆展示、志愿活动、社区共享规则。"),
            ("形式", "“空间要温暖、可识别、有主题，但不能华而不实”。", "流线固定、趣味性不足；节点记忆弱；家具和材料适老性不足。", "主题花园、色彩分区、触觉材料、半围合座椅、清晰边界。"),
        ],
        [0.85, 1.9, 1.9, 1.75],
        header_fill="F2F4F7",
        size=7.0,
    )
    h2(doc, "4.3 工作会议成果矩阵")
    add_table(
        doc,
        ["价值", "目标", "事实信息", "需求", "策划概念"],
        [
            ("环境", "营造疗愈景观；提升日照与树荫停留品质；强化空间识别。", "植被维护粗放；导视不清；南向和中庭停留频繁；街心公园活力不足。", "连续绿荫、冬季日照、芳香/声景、休息座椅、低维护植物。", "五感疗愈花园、翠阴廊、凝神之园、芳香与声景节点。"),
            ("功能", "满足生活照护、康复运动、家属探访和社区共享。", "散步、晒太阳、交谈、锻炼、探访高频；互动设施不足。", "地滚球、园艺课程、探访座椅、活动展示、护理陪同。", "康竞园、丰收径、共聚之园、八方园、家属探访节点。"),
            ("交通", "形成安全连续、可识别、可休息的慢行系统。", "扶手不连续；台阶/高差阻隔；消防口打断；路径识别弱。", "无障碍坡道、连续扶手、防滑铺装、约50米休息点、夜间照明。", "探索康养径、内外分流、图形导视、节点打卡。"),
            ("经济", "低成本、短周期、可分期、可维护。", "存量场地更新；施工影响敏感；维护力量有限。", "针灸式微更新、模块化设施、可替换展陈、低维护种植。", "近期安全底线包、中期节点激活包、长期运营维护包。"),
            ("文化", "降低机构封闭感；增强记忆疗愈与社区连接。", "记忆展陈不足；认知友好宣传弱；家属与志愿者参与场景不足。", "老物件/老照片、认知症友好宣教、节庆展示、社区共建。", "忆录园、认知体感墙、社区共享界面、志愿活动花园。"),
            ("形式", "形成适老尺度、温暖美观、易识别的空间形象。", "流线固定、趣味性不足；节点记忆弱；家具与材料适老性不足。", "主题化节点、色彩分区、触觉材料、半围合座椅、明确边界。", "五园一廊一径；街心公园南北两区；可触摸、可参与景观。"),
        ],
        [0.65, 1.45, 1.55, 1.45, 1.45],
        header_fill="DDEBF7",
        size=6.7,
    )

    h1(doc, "5. 概念性策划")
    h2(doc, "5.1 总体概念")
    para(doc, "总体概念为“探索康养径串联的疗愈型适老化户外生活网络”。该概念强调三个层次：第一，以连续康养径解决安全可达；第二，以五园一廊一径承载院内日常生活；第三，以街心公园南北分区形成对外展示、社区共享和运动干预界面。")
    h2(doc, "5.2 空间结构")
    add_table(
        doc,
        ["层级", "组成", "核心功能", "控制要点"],
        [
            ("一径", "探索康养径。", "连续步行、康复训练、节点串联、展示参观。", "闭环或准闭环；按无障碍标准复核坡度、净宽、扶手和防滑。"),
            ("一廊", "翠阴廊。", "阴凉通行、植物疗愈、慢行停留。", "攀藤植物、遮阴、可识别导视、适老座椅。"),
            ("五园", "八方园、三松园、隐逸园、康竞园、忆录园，并结合丰收径园艺系统。", "入口聚集、活动展示、安静休憩、低强度运动、记忆疗愈、园艺课程。", "每园形成明确主题、借力座椅、护理可视范围和运营活动。"),
            ("街心公园", "南向认知互动与宣教区；北向运动干预与社交活动区。", "认知体感、科普展示、地滚球、平衡训练、社区共聚。", "处理高差、消防口打断和可管理开放边界。"),
        ],
        [0.75, 1.55, 1.9, 2.25],
        size=7.2,
    )
    h2(doc, "5.3 核心策略")
    add_table(
        doc,
        ["策略", "策划含义", "空间动作", "管理动作"],
        [
            ("路径缝合", "把碎片空间串联为连续、安全、有目的地的康养路径。", "环线、节点、休息点、导视、内外分流。", "试走评估、风险点巡检。"),
            ("针灸式微更新", "先更新高频和高风险节点，以点带面。", "模块化家具、可替换展陈、局部铺装与扶手补齐。", "分期施工、使用反馈迭代。"),
            ("多感官疗愈", "从“看景”转向“参与景观”。", "芳香、声景、触觉、园艺、季相植物。", "园艺课、情绪观察、康复活动记录。"),
            ("场景复合", "让散步、日照、社交、康复和探访在同一网络中发生。", "共聚座椅、探访节点、运动花园、记忆展陈。", "活动日历、社工和志愿者联动。"),
            ("运营闭环", "把设计成果转化为长期可使用的服务。", "维护储物、展陈接口、活动电源/水源预留。", "维护制度、责任人、年度评估。"),
        ],
        [0.9, 1.75, 1.85, 1.95],
        size=7.3,
    )
    h2(doc, "5.4 设计控制指标")
    bullets(doc, [
        "慢行路径：优先形成闭环或准闭环，避免老人走入无效尽端；关键节点设置清晰方向标识。",
        "休息系统：结合日照、树荫和护理视线布置座椅，约50米设置可停靠节点，座椅应便于起身借力。",
        "植物系统：优先低维护、无明显安全风险、季相稳定的适生植物；芳香植物需避免过强刺激和过敏风险。",
        "认知友好：导视采用图形、颜色和节点命名组合，文字信息避免过密。",
        "运营接口：每类空间均对应至少一类可持续活动，如园艺、康复、节庆、志愿陪伴或认知宣教。",
    ])

    h1(doc, "HECTTEAS 矩阵分析")
    h2(doc, "5.5 价值排序")
    para(doc, "参考图3，HECTTEAS 不是固定答案，而是针对特定项目的价值排序工具。本项目以养老护理院更新为对象，安全和人文因素应置于最高优先级；环境、技术和运营时间是支撑条件；经济、美学、文化则需服务于长期使用和尊严生活。")
    add_table(
        doc,
        ["排序", "价值因素", "优先级理由", "设计含义"],
        [
            ("1", "Safety 安全因素", "跌倒、迷失、夜间风险和消防边界直接决定老人能否使用。", "先补齐扶手、防滑、照明、坡道和求助系统。"),
            ("2", "Human 人文因素", "老人差异化能力、尊严、参与感是养老空间的核心。", "空间支持自理、介助、介护老人共同使用。"),
            ("3", "Environmental 环境因素", "自然疗愈和热舒适影响户外活动意愿。", "组织树荫、日照、芳香、声景和低维护植物。"),
            ("4", "Technical 技术因素", "无障碍、照明、导视和运维系统决定方案能否落地。", "建立可审查、可巡检的技术清单。"),
            ("5", "Temporal 时间因素", "既有机构施工需分期，运营也需逐步导入。", "近期安全、中期激活、长期机制固化。"),
            ("6", "Cultural 文化因素", "记忆疗愈和社区连接提升机构温度。", "设置忆录园、认知宣教和社区共享界面。"),
            ("7", "Economic 经济因素", "长期维护成本影响方案生命力。", "采用模块化、低维护、可替换策略。"),
            ("8", "Aesthetic 美学因素", "美学应服务识别、舒适和尊严，而非单纯装饰。", "形成温暖、清晰、主题化的适老空间。"),
        ],
        [0.45, 1.25, 2.6, 2.1],
        header_fill="E2F0D9",
        size=7.3,
        align_first=True,
    )
    h2(doc, "5.6 HECTTEAS 专业矩阵")
    add_table(
        doc,
        ["价值", "目标", "事实", "需求", "问题", "概念"],
        [
            ("人文\nHUMAN", "从被动照护转向积极参与。", "散步、日照、交谈、探访高频；老人能力差异明显。", "可达、可停、可看护、可交往。", "空间支持不足会降低外出意愿。", "康养径、探访花园、护理观察点。"),
            ("环境\nENVIRONMENTAL", "构建自然疗愈和热舒适环境。", "植物层次不足，日照与树荫分布不均。", "遮阴、冬季日照、芳香、声景。", "只做绿化无法支撑持续使用。", "五感花园、翠阴廊、芳香节点。"),
            ("文化\nCULTURAL", "增强记忆疗愈和社区连接。", "记忆展陈与认知友好宣传不足。", "老照片、老物件、节庆、志愿活动。", "机构封闭感强，社区参与场景不足。", "忆录园、认知体感墙、共享界面。"),
            ("技术\nTECHNICAL", "落实无障碍、导视、照明和求助。", "扶手不连续，高差与消防口打断路径。", "坡道、扶手、防滑、照明、求助。", "技术不连续直接形成安全风险。", "安全设施包、色彩导视、夜间照明。"),
            ("时间\nTEMPORAL", "分期实施，减少运营扰动。", "既有养老机构不能长期中断运营。", "近期补短板，中期激活，长期维护。", "一次性大改造难以根据反馈迭代。", "0-6月、6-18月、18月以上分期。"),
            ("经济\nECONOMIC", "控制造价和长期维护压力。", "维护人力有限，设施需长期完好。", "模块化、低维护、可替换。", "高维护景观会成为院方负担。", "针灸式微更新、标准化采购。"),
            ("美学\nAESTHETIC", "形成温暖、清晰、有尊严的形象。", "现状趣味性和节点记忆不足。", "主题花园、季相色彩、触觉材料。", "装饰化美学容易脱离老人行为。", "五园一廊一径主题系统。"),
            ("安全\nSAFETY", "把跌倒、迷失、夜间和消防风险作为底线。", "高差、台阶、扶手断点、照明不足。", "连续扶手、防滑、休息点、求助。", "安全不稳，疗愈和社交无法成立。", "安全闭环路径、风险巡检表。"),
        ],
        [0.72, 1.08, 1.17, 1.1, 1.12, 1.16],
        header_fill="D9EAD3",
        size=6.0,
    )

    h1(doc, "6. 空间分配表")
    para(doc, "空间分配采用“比例控制+面积公式”。A 为施工图阶段经 CAD/测绘复核后的总更新面积。为避免在缺少测绘资料时制造虚假精确值，表中同时给出每1000平方米更新面积的折算指标。")
    add_table(
        doc,
        ["空间类型", "比例", "面积公式", "每1000㎡折算", "主要内容"],
        [
            ("连续康养径与无障碍慢行", "28%", "0.28A", "280㎡", "环线、扶手、防滑、坡道、导视、休息点。"),
            ("日照与安静休憩", "14%", "0.14A", "140㎡", "半围合座椅、可移动家具、遮阴与冬季日照。"),
            ("多感官疗愈花园", "18%", "0.18A", "180㎡", "季相、触觉、声景、芳香植物、冥想停留。"),
            ("社交探访与展示活动", "12%", "0.12A", "120㎡", "探访座椅、活动展示、节庆布置、社区宣教。"),
            ("康复运动与地滚球", "10%", "0.10A", "100㎡", "环形跑道、地滚球、平衡训练、护理陪同。"),
            ("园艺疗法与共享种植", "8%", "0.08A", "80㎡", "抬高种植床、采摘工具、可食植物、园艺课程。"),
            ("导视照明求助与管理设施", "6%", "0.06A", "60㎡", "色彩导视、图形标识、照明、紧急求助、维护储物。"),
            ("维护缓冲与弹性机动", "4%", "0.04A", "40㎡", "消防边界、可替换展陈、临时活动接口。"),
        ],
        [1.55, 0.55, 0.75, 0.85, 2.75],
        size=7.0,
    )
    add_table(
        doc,
        ["分区", "建议占比", "角色", "重点节点"],
        [
            ("院内庭院更新区", "约60%", "日常高频使用、护理可控、低风险康复、安静疗愈。", "探索康养径、八方园、三松园、隐逸园、翠阴廊、丰收径、康竞园、忆录园。"),
            ("街心公园南向区", "约18%", "认知症互动、宣教展示、对外形象、低复杂度停留。", "认知体感墙、老照片展陈、触觉材料、多感官植物。"),
            ("街心公园北向区", "约22%", "运动干预、地滚球、社区共聚、外出活动延展。", "轨迹之园、合奏之园、平衡训练、声景互动。"),
        ],
        [1.45, 0.85, 2.1, 2.05],
        header_fill="F2F4F7",
        size=7.3,
    )

    h1(doc, "7. 空间图表")
    h2(doc, "7.1 行为-空间-运营转译图表")
    add_table(
        doc,
        ["行为", "空间需求", "设施落位", "运营活动"],
        [
            ("散步", "连续、可识别、可休息的慢行环线。", "探索康养径、扶手、遮阴、休息节点、里程打卡。", "晨间散步、饭后慢行、康复步行记录。"),
            ("晒太阳", "南向、半开放、可移动座椅和冬季日照。", "中庭、南向停留点、可移动椅、暖色铺装。", "午间日照陪伴、家属探访、节气活动。"),
            ("交谈", "小尺度围合、可面对面、护理人员可观察。", "八方园、三松园、共聚之园、家属座椅。", "茶话会、生日会、志愿陪伴。"),
            ("锻炼", "低强度、低冲击、可陪同与可中止。", "康竞园、地滚球、平衡训练、环形跑道。", "康复训练、地滚球小组、健康观察记录。"),
            ("认知刺激", "低复杂度、可触摸、可回忆、可重复。", "认知体感墙、老物件展示、色彩导视、芳香植物。", "记忆叙事、认知症友好宣教。"),
            ("园艺参与", "可站可坐、便于护理陪同、工具安全。", "丰收径、抬高种植床、可食植物、采摘工具。", "园艺课程、采摘日、植物认养。"),
        ],
        [0.8, 1.85, 1.95, 1.85],
        size=7.0,
    )
    h2(doc, "7.2 空间网络图表")
    add_table(
        doc,
        ["入口/展示", "日常路径", "安静疗愈", "参与活动", "外向共享"],
        [
            ("八方园\n门面聚集\n树下座椅", "探索康养径\n串联院内外", "隐逸园\n安静休憩", "丰收径\n园艺采摘", "街心南向\n认知宣教"),
            ("三松园\n灵活家具\n方院活动", "翠阴廊\n阴凉通行", "凝神之园\n五感安抚", "康竞园\n平衡训练", "街心北向\n运动共聚"),
            ("忆录园\n老物件展示", "导视系统\n色彩引导", "芳香/声景节点", "共聚之园\n家属探访", "社区界面\n可管理开放"),
        ],
        [1.28, 1.28, 1.28, 1.28, 1.28],
        header_fill="FCE4D6",
        size=6.8,
    )
    h2(doc, "7.3 分期实施图表")
    add_table(
        doc,
        ["阶段", "目标", "工程重点", "运营导入", "验收重点"],
        [
            ("近期\n0-6个月", "补齐安全底线。", "扶手、防滑、照明、座椅、导视、坡道、消防口边界。", "护理陪同路线、老人试走、跌倒风险复盘。", "无障碍通行、夜间安全、休息点连续性。"),
            ("中期\n6-18个月", "激活核心节点。", "认知体感墙、园艺疗法、地滚球、声景与芳香节点。", "园艺课程、康复小组、家属探访活动。", "参与人数、停留时长、设施完好率。"),
            ("长期\n18个月以上", "形成机制闭环。", "可替换展陈、植物更新、社区接口、活动储物与维护系统。", "志愿者、学校、社区、医养结合记录。", "年度运营日历、满意度、维护成本。"),
        ],
        [0.85, 1.0, 1.75, 1.45, 1.4],
        header_fill="E2F0D9",
        size=6.9,
    )

    h1(doc, "8. 总结")
    para(doc, "深圳市养老护理院景观改造的关键，不在于增加装饰性绿化，而在于把老人真实、重复、脆弱但有活力的日常行为转化为空间系统。院内通过“五园一廊一径”形成安全可控的日常生活网络，街心公园通过南北分区形成认知友好、运动干预和社区共享界面。")
    para(doc, "本报告的深化结论包括：第一，安全性是所有目标的前置条件，应优先形成连续扶手、防滑铺装、夜间照明、休息点和紧急求助系统；第二，疗愈性不应停留在观赏植物，而应通过园艺、声景、芳香、触觉和康复活动形成可参与的日常机制；第三，社交性应通过半围合座椅、家属探访节点、社区共享界面和志愿活动被空间化；第四，认知友好应通过图形导视、色彩分区、记忆节点和低复杂度路径降低老人焦虑。")
    para(doc, "最终目标是实现三重转变：从“养老空间”向“生活空间”的转变，从“被动照护”向“积极参与”的转变，从“封闭机构”向“健康社区”的转变。")

    h1(doc, "9. 建筑策划任务书")
    para(doc, "本任务书作为后续方案设计、初步设计和施工图深化的控制依据。面积以 A 表示总更新面积，待测绘/CAD复核后换算为平方米；每1000平方米折算值用于课堂汇报和策划阶段快速核对。")
    add_table(
        doc,
        ["功能分区", "内容", "面积", "要求"],
        [
            ("A区：院内探索康养径", "连续慢行环线、扶手、防滑铺装、休息节点、导视、照明、护理观察点。", "0.28A\n每1000㎡为280㎡", "按无障碍要求复核；约50米设置休息点；路径避免无效折返和护理盲区。"),
            ("B区：院内五园节点", "八方园、三松园、隐逸园、康竞园、忆录园及节点家具。", "0.24A\n每1000㎡为240㎡", "每园形成主题、停留方式和运营活动；座椅便于借力、轮椅停靠和护理陪同。"),
            ("C区：翠阴廊与多感官疗愈系统", "绿廊、芳香植物、触觉材料、声景装置、季相植物、冥想停留。", "0.18A\n每1000㎡为180㎡", "植物低维护、无明显毒刺风险；声景控制音量；兼顾遮阴与通风。"),
            ("D区：丰收径与园艺疗法", "抬高种植床、可食植物、采摘工具、园艺课程接口、储物设施。", "0.08A\n每1000㎡为80㎡", "种植床满足坐姿/站姿操作；工具安全可控；课程与护理、社工排期联动。"),
            ("E区：街心公园南向认知互动与宣教", "认知体感墙、老照片/老物件展陈、科普互动、低复杂度导视。", "0.10A\n每1000㎡为100㎡", "内容避免信息过载；图形、色彩和文字高度适老；兼顾对外展示与院方管理。"),
            ("F区：街心公园北向运动干预与社交", "地滚球、平衡训练、合奏之园、社区共聚座椅、活动边界。", "0.12A\n每1000㎡为120㎡", "运动设施低冲击、可陪同、可中止；处理消防口和高差；避免与城市通行冲突。"),
            ("G区：导视照明求助与运维系统", "色彩导视、图形标识、夜间照明、紧急求助、维护储物、可替换展陈。", "0.06A\n每1000㎡为60㎡", "建立点位清单和巡检制度；照明、求助、消防、物业维护同步审查。"),
            ("H区：弹性与缓冲空间", "消防口边界、设备缓冲、临时活动接口、施工转换区。", "0.04A\n每1000㎡为40㎡", "不得侵占消防和必要通行；支持节庆活动、志愿活动和后期设施迭代。"),
        ],
        [1.45, 2.15, 1.05, 1.8],
        header_fill="E2F0D9",
        size=6.8,
    )
    h2(doc, "9.1 后续设计任务")
    numbered(doc, [
        "补充测绘资料：CAD红线、竖向标高、树木现状、消防口、排水、照明和设备点位。",
        "开展专项复核：无障碍坡度、通行净宽、扶手连续性、防滑等级、夜间照度、护理视线。",
        "深化节点方案：对八方园、三松园、隐逸园、翠阴廊、丰收径、康竞园、忆录园和街心公园南北区形成节点详图。",
        "建立运营手册：明确活动日历、责任部门、维护频次、志愿者机制、设施巡检和老人反馈渠道。",
        "形成评估闭环：试运行后按参与人数、停留时长、跌倒/险情记录、家属满意度和维护成本进行复盘。",
    ])

    add_references(doc)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
