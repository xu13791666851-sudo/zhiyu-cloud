from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


OUT = Path(__file__).resolve().parent / "深圳市养老护理院景观改造建筑策划报告.docx"
SOURCE_PPT = r"C:\Users\19727\Desktop\建筑策划\建筑策划汇报.pptx"


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, name="Microsoft YaHei", size=10.5, color="1F2937"):
    for run in paragraph.runs:
        set_run_font(run, name=name, size=size, color=color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color="1F2937", size=9.2, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths, header_fill="E8EEF5", font_size=9.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for idx, h in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], header_fill)
        set_cell_text(table.rows[0].cells[idx], h, bold=True, color="0B2545", size=font_size)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            set_cell_text(row.cells[idx], str(value), size=font_size)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color="1F2937")


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, size=10.5, color="1F2937")


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_font(p, size=16, color="2E74B5")


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    set_paragraph_font(p, size=13, color="2E74B5")


def para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=10.5, bold=True, color="0B2545")
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=10.5, color="1F2937")
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color="1F2937")
    return p


def set_section_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)


def set_section_portrait(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        st = styles[style_name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("深圳市养老护理院景观改造建筑策划报告")
    set_run_font(hr, size=8, color="6B7280")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("基于积极老龄化理念的疗愈型适老化景观更新研究")
    set_run_font(fr, size=8, color="6B7280")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("深圳市养老护理院景观改造建筑策划报告")
    set_run_font(r, size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    r = subtitle.add_run("基于积极老龄化理念的疗愈型适老化景观更新研究")
    set_run_font(r, size=13, color="2E74B5")

    meta = [
        ("报告性质", "建筑策划/景观更新策划报告"),
        ("研究对象", "深圳市养老护理院院内庭院、街心公园、室外慢行系统与集中活动区域"),
        ("依据资料", f"源PPT：{SOURCE_PPT}；公开政策、标准与循证研究文献"),
        ("编制日期", "2026年6月8日"),
    ]
    add_table(doc, ["项目", "内容"], meta, [1.5, 5.0], header_fill="F2F4F7", font_size=9.2)

    para(
        doc,
        "核心判断：本项目不是单点景观美化，而是以“安全底线、疗愈促进、社交参与、认知友好”为主线，把院内庭院与街心公园组织为连续、可停留、可参与、可运营的户外康养生活系统。",
        bold_lead="核心判断：",
    )
    doc.add_page_break()

    h1(doc, "1. 选址")
    para(
        doc,
        "项目位于深圳市南山区桃源街道龙珠七路与龙苑路交界处的深圳市养老护理院及其相邻街心公园界面。根据公开项目资料，深圳市养老护理院用地面积约1万平方米，建筑面积约4万平方米，设置床位约800张，是以长期照护、护理康复和医养结合为主要功能的养老机构。"
    )
    h2(doc, "1.1 选址性质")
    add_bullets(
        doc,
        [
            "既有机构更新：本策划以既有养老护理院及周边绿地更新为主，重点不是重新购地，而是在高密度城区中释放存量场地的康养价值。",
            "院内外双界面：院内庭院承接日常护理、低强度活动与安静休憩；街心公园承接外出活动、社区共享、认知症友好宣教和养老院形象展示。",
            "慢行可达优先：场地选择的价值在于可形成连续慢行系统，并把老人、家属、护理人员、志愿者和社区居民的流线进行温和分流。",
        ],
    )
    h2(doc, "1.2 选址评价")
    add_table(
        doc,
        ["评价维度", "现状判断", "策划要求"],
        [
            ("可达性", "机构位于城市居住片区内，具备家属探访和社区参与基础。", "强化入口识别、无障碍坡道、连续扶手与路口提示，降低外出门槛。"),
            ("安全性", "院内空间相对可控，但街心公园存在高差、消防口打断和线性空间不连续。", "以“院内安全闭环+街心公园可看护开放”为原则分级组织。"),
            ("景观潜力", "现状有庭院、树荫与街心绿地资源，但植物层次、互动性和疗愈属性不足。", "导入五感花园、园艺疗法、声景和认知互动设施。"),
            ("运营潜力", "已形成散步、晒太阳、交谈、简单锻炼和探访等高频行为。", "把高频行为转译为固定活动节点，并形成护理、康复、社工、志愿者共同运营机制。"),
        ],
        [1.15, 2.55, 2.8],
    )

    h1(doc, "2. 项目背景研究")
    para(
        doc,
        "我国人口老龄化持续加深，养老机构正在从“床位供给”和“被动照护”转向“健康促进、长期参与和尊严生活”。积极老龄化框架强调健康、参与和保障；年龄友好城市框架强调户外空间、交通、社会参与、尊重与包容等要素。对养老护理院而言，户外空间不再只是绿化附属，而是康复训练、情绪调节、认知刺激、家属探访和公共生活的基础设施。"
    )
    h2(doc, "2.1 项目更新动因")
    add_bullets(
        doc,
        [
            "需求侧：入住老人存在稳定的散步、日照、社交、探访和低强度锻炼需求；半失能、介助老人需要更连续的扶手、休息点和可视化看护。",
            "空间侧：现状存在植被维护粗放、导视不足、扶手不连续、互动装置弱、集中活动场地日照和座椅不足等问题。",
            "运营侧：现有空间缺少可持续组织活动的场景和设施接口，护理、康复、社工和社区志愿活动难以沉淀为日常机制。",
            "城市侧：街心公园可从边角绿地转为养老院外向展示、社区共享和长者外出活动延展的康养界面。",
        ],
    )
    h2(doc, "2.2 策划目标体系")
    add_table(
        doc,
        ["目标", "内涵", "落位方式"],
        [
            ("安全性", "连续慢行、扶手、防滑、照明、紧急求助、可视化看护。", "康养环线、50米左右休息节点、坡道与台阶协同、夜间照明。"),
            ("疗愈性", "自然介入、园艺疗法、多感官刺激、冥想休憩和康复活动。", "感官花园、芳香植物、声景装置、抬高种植床、安静花园。"),
            ("社交性", "老人、家属、护理人员、志愿者和社区之间的低压力交往。", "半围合座椅、探访节点、群体活动场、共享种植和节庆展示。"),
            ("认知友好", "降低迷失、焦虑与路径判断压力，强化识别和记忆。", "图形导视、色彩分区、记忆展陈、认知体感墙和低复杂度路径。"),
        ],
        [1.0, 2.65, 2.85],
    )

    h1(doc, "3. 收集与分析信息")
    para(
        doc,
        "信息收集分为四类：源PPT现场照片与观察结论、场地空间条件、老人行为需求、政策标准与循证研究。由于本次未获得CAD红线、精确面积和日照模拟文件，空间分配以“更新范围100%”进行比例化策划，后续应由测绘/CAD复核转化为平方米指标。"
    )
    h2(doc, "3.1 信息清单")
    add_table(
        doc,
        ["信息类型", "已获得内容", "分析用途", "后续补充"],
        [
            ("场地现状", "院内庭院、街心公园、慢行系统、集中活动区域照片与问题归纳。", "识别安全缺口、体验缺口、认知缺口和运营缺口。", "CAD、竖向标高、无障碍坡度、铺装防滑参数。"),
            ("行为观察", "散步、晒太阳、交谈、锻炼、探访等高频行为及时间规律。", "形成行为-空间-设施转译表。", "分时段人流计数、停留热力、护理陪同记录。"),
            ("使用人群", "自理、介助、介护老人及家属、护理人员、志愿者。", "决定空间开放层级、看护距离和活动强度。", "认知症比例、行动能力分级、跌倒风险点位。"),
            ("政策与研究", "积极老龄化、年龄友好、疗愈景观、无障碍与养老设施标准。", "为策划目标和控制指标提供依据。", "地方审批条件、消防与运维部门意见。"),
        ],
        [1.05, 2.0, 1.9, 1.55],
        font_size=8.6,
    )
    h2(doc, "3.2 问题综合")
    add_table(
        doc,
        ["问题", "表现", "空间后果", "更新方向"],
        [
            ("安全缺口", "扶手、坡道、座椅借力点、夜间照明未形成连续系统。", "老人外出意愿降低，护理陪同压力增大。", "补齐连续扶手、防滑铺装、照明、紧急求助与可视看护点。"),
            ("体验缺口", "景观以观看为主，触摸、声音、园艺和运动参与不足。", "场地停留时间短，疗愈效益弱。", "导入五感花园、园艺课程、地滚球和平衡训练。"),
            ("认知缺口", "导视、图形识别、节点记忆和色彩引导不足。", "方向感下降老人易焦虑，路径选择负担高。", "建立色彩分区、低复杂度导视和记忆节点。"),
            ("运营缺口", "缺少可固定组织活动的场景与维护机制。", "改造后容易回到“只看不参与”。", "以活动日历、志愿机制、护理记录和模块化维护闭环保障持续使用。"),
        ],
        [1.05, 1.8, 1.75, 1.9],
        font_size=8.6,
    )

    h1(doc, "4. 工作会议")
    para(
        doc,
        "建筑策划应通过多方工作会议把“设计意图”转化为“可建、可管、可运营”的任务书。建议设置五轮会议，并在每轮形成纪要、决策清单和下一步责任人。"
    )
    add_table(
        doc,
        ["会议", "参会方", "核心议题", "输出成果"],
        [
            ("启动会", "院方、护理部、康复师、设计团队、物业运维", "明确范围、边界、目标人群、施工影响与资料清单。", "项目任务书、调研计划、风险清单。"),
            ("现状诊断会", "院方、护理人员、老人代表、家属代表", "核对高频路径、跌倒风险点、日照停留点、噪声与护理盲区。", "问题地图、行为热力、需求优先级。"),
            ("概念策划会", "院方、设计团队、运营/社工、康复师", "确认四位一体目标、五园一廊一径、街心公园南北分区。", "概念方案、功能框架、空间分配初稿。"),
            ("技术协调会", "设计、消防、无障碍、机电照明、景观施工、运维", "坡道、扶手、铺装、照明、排水、消防口、植物维护和设施安全。", "技术控制清单、材料样板、造价边界。"),
            ("运营落地会", "护理部、康复师、社工、志愿者组织、社区代表", "晨间活动、园艺课程、地滚球、认知友好宣教、社区共享规则。", "年度运营日历、维护制度、试运行评估表。"),
        ],
        [0.85, 1.65, 2.2, 1.8],
        font_size=8.3,
    )

    h1(doc, "5. 概念性策划")
    h2(doc, "5.1 总体概念")
    para(
        doc,
        "总体概念为“探索康养径串联的疗愈型适老化户外生活网络”。院内以“五园一廊一径”承接日常生活、康复和记忆疗愈；街心公园以南向认知互动与宣教区、北向运动干预与社交活动区承接外出活动和社区共享。"
    )
    h2(doc, "5.2 空间结构")
    add_table(
        doc,
        ["空间层级", "组成", "核心功能", "控制要点"],
        [
            ("一径", "探索康养径", "连续步行、康复训练、节点串联、展示参观。", "闭环或准闭环；50米左右休息点；连续扶手与防滑铺装。"),
            ("一廊", "翠阴廊", "阴凉通行、植物疗愈、慢行停留。", "攀藤植物、遮阴、可识别导视、无障碍通行净宽。"),
            ("五园", "八方园、三松园、隐逸园、康竞园、忆录园，并结合丰收径园艺系统", "入口聚集、活动展示、安静休憩、低强度运动、记忆疗愈、园艺课程。", "每个节点有明确主题、座椅借力点和护理可视范围。"),
            ("街心公园", "南向认知互动与宣教区、北向运动干预与社交活动区", "认知体感、科普展示、地滚球、平衡训练、社区共聚。", "处理高差、消防口打断和开放管理边界。"),
        ],
        [1.0, 1.6, 2.05, 1.85],
        font_size=8.5,
    )
    h2(doc, "5.3 策略包")
    add_numbered(
        doc,
        [
            "路径缝合：把碎片空间串联为连续康养径，形成可达、可停、可看护的户外活动骨架。",
            "针灸式微更新：优先更新高频节点，以模块化设施、可替换展陈和低维护种植降低施工与运营压力。",
            "多感官疗愈系统：将视觉、触觉、听觉、嗅觉、味觉和运动觉落到植物、铺装、声景、园艺和康复活动中。",
            "场景复合：把散步、社交、宣教、园艺、康复运动和家属探访叠合在日常路径上。",
            "运营闭环：近期补安全底线，中期激活核心节点，长期形成活动排期、社区共享、医养联动和维护评估制度。",
        ],
    )

    h1(doc, "HECTTEAS 矩阵分析")
    para(
        doc,
        "本报告将 HECTTEAS 定义为八个策划评价维度：H-Health/Healing 健康疗愈，E-Environment 环境生态，C-Culture/Community 文化社区，T-Technology 技术系统，T-Time 时间阶段，E-Economy 经济运营，A-Aesthetics 美学体验，S-Safety 安全韧性。矩阵用于把概念目标转化为可检查的设计和运营要求。"
    )
    add_table(
        doc,
        ["维度", "现状风险", "策划动作", "评价指标"],
        [
            ("H 健康疗愈", "活动刺激单一，景观疗愈属性弱。", "园艺疗法、冥想花园、芳香植物、地滚球与低强度康复训练。", "每日活动时长、参与人数、情绪/睡眠观察、康复记录。"),
            ("E 环境生态", "植物层次不足，维护压力和季相体验不均衡。", "低维护乡土/适生植物、遮阴、雨水友好铺装、可食植物。", "遮阴覆盖、维护频次、植物成活率、热舒适反馈。"),
            ("C 文化社区", "机构封闭感强，记忆与社区连接不足。", "老物件展示、认知症友好宣教、家属探访节点、志愿活动。", "活动次数、家属停留、社区共建项目、老人记忆叙事采集。"),
            ("T 技术系统", "导视、照明、求助和看护点不连续。", "图形导视、色彩分区、照明分级、紧急求助、可视化看护。", "导视识别率、夜间照度达标率、求助响应时间。"),
            ("T 时间阶段", "施工影响与运营导入缺少分期。", "近期安全补短板；中期节点激活；长期机制固化。", "分期完成率、施工扰动天数、试运行反馈闭环。"),
            ("E 经济运营", "一次性景观投入易变成后期维护负担。", "模块化设施、可替换展陈、低维护种植、活动赞助和志愿机制。", "单位面积维护成本、设施完好率、活动运营成本。"),
            ("A 美学体验", "景观趣味和可识别性不足。", "五园一廊一径主题化、季相色彩、声景与触感节点。", "停留时长、拍照/展示频次、老人和家属满意度。"),
            ("S 安全韧性", "高差、台阶、扶手断点、地面防滑和消防口打断。", "坡道与台阶协同、连续扶手、防滑铺装、夜间照明、消防口清晰边界。", "跌倒/险情记录、无障碍通行率、消防与运维复核结果。"),
        ],
        [1.0, 1.65, 2.55, 1.3],
        font_size=8.0,
    )

    h1(doc, "6. 空间分配表")
    para(
        doc,
        "以下为空间策划阶段比例表，适用于缺少精确测绘面积时的任务书控制。后续可将总更新面积A代入，得到各类空间面积；例如某类空间面积= A × 建议比例。"
    )
    add_table(
        doc,
        ["空间类型", "建议比例", "主要位置", "设施与内容", "优先级"],
        [
            ("连续康养径与无障碍慢行", "28%", "院内路径、街心公园衔接段", "环线、扶手、防滑铺装、坡道、导视、50米左右休息点", "高"),
            ("日照与安静休憩", "14%", "中庭、南向停留点、隐逸园", "半围合座椅、可移动家具、遮阴与冬季日照平衡", "高"),
            ("多感官疗愈花园", "18%", "凝神之园、共聚之园、翠阴廊、芳香节点", "视觉季相、触觉材料、声景、芳香植物、冥想停留", "高"),
            ("社交探访与展示活动", "12%", "八方园、三松园、入口与街心公园界面", "家属探访座椅、活动展示、节庆布置、社区宣教", "中高"),
            ("康复运动与地滚球", "10%", "康竞园、街心公园北向区域", "环形跑道、地滚球、平衡训练、护理陪同看护点", "中高"),
            ("园艺疗法与共享种植", "8%", "丰收径、共享种植节点", "抬高种植床、采摘工具、可食植物、园艺课程", "中"),
            ("导视照明求助与管理设施", "6%", "全场关键节点", "色彩导视、图形标识、照明、紧急求助、维护储物", "高"),
            ("维护缓冲与弹性机动", "4%", "边界、消防口、设备与临时活动区", "消防通道控制、可替换展陈、临时活动接口", "中"),
        ],
        [1.45, 0.7, 1.45, 2.1, 0.55],
        font_size=7.8,
    )
    add_table(
        doc,
        ["分区", "建议占比", "承担角色", "重点节点"],
        [
            ("院内庭院更新区", "约60%", "老人日常高频使用、护理可控、低风险康复、安静疗愈。", "探索康养径、八方园、三松园、隐逸园、翠阴廊、丰收径、康竞园、忆录园。"),
            ("街心公园南向区", "约18%", "认知症互动、宣教展示、对外形象、低复杂度停留。", "认知体感墙、老照片展陈、触觉材料、多感官植物。"),
            ("街心公园北向区", "约22%", "运动干预、地滚球、社区共聚、外出活动延展。", "轨迹之园、合奏之园、平衡训练、声景互动。"),
        ],
        [1.35, 0.8, 2.25, 2.1],
        font_size=8.2,
    )

    h1(doc, "7. 空间图表")
    h2(doc, "7.1 行为-空间-运营转译图表")
    add_table(
        doc,
        ["高频行为", "空间需求", "设施落位", "运营活动"],
        [
            ("散步", "连续、可识别、可休息的慢行环线。", "探索康养径、扶手、遮阴、休息节点、里程打卡。", "晨间散步、饭后慢行、康复步行记录。"),
            ("晒太阳", "南向、半开放、可移动座椅和冬季日照。", "中庭、南向停留点、可移动椅、暖色铺装。", "午间日照陪伴、家属探访、节气活动。"),
            ("交谈", "小尺度围合、可面对面、护理人员可观察。", "八方园、三松园、共聚之园、家属座椅。", "茶话会、生日会、志愿陪伴。"),
            ("锻炼", "低强度、低冲击、可陪同与可中止。", "康竞园、地滚球、平衡训练、环形跑道。", "康复训练、地滚球小组、健康观察记录。"),
            ("认知刺激", "低复杂度、可触摸、可回忆、可重复。", "认知体感墙、老物件展示、色彩导视、芳香植物。", "记忆叙事、认知症友好宣教、家属共创展陈。"),
            ("园艺参与", "可站可坐、便于护理陪同、工具安全。", "丰收径、抬高种植床、可食植物、采摘工具。", "园艺课程、采摘日、植物认养。"),
        ],
        [1.0, 1.8, 1.85, 1.85],
        font_size=8.2,
    )
    h2(doc, "7.2 空间网络图表")
    add_table(
        doc,
        ["入口/展示", "日常康养路径", "安静疗愈", "参与活动", "外向共享"],
        [
            ("八方园\n门面聚集、树下座椅、活动展示", "探索康养径\n串联院内节点与街心公园", "隐逸园\n绿篱隔离、安静休憩", "丰收径\n种植箱、园艺课程、采摘互动", "街心公园南向\n认知互动与宣教"),
            ("三松园\n灵活家具、方院活动、记忆节点", "翠阴廊\n拱形绿廊、阴凉通行", "凝神之园\n五感安抚、冥想停留", "康竞园\n环形跑道、地滚球、平衡训练", "街心公园北向\n运动干预与社区共聚"),
            ("忆录园\n老物件展示、记忆疗愈", "导视系统\n图形识别、色彩引导", "芳香/声景节点\n嗅觉与听觉疗愈", "共聚之园\n家属探访、志愿活动", "社区界面\n可管理开放、养老形象展示"),
        ],
        [1.3, 1.3, 1.3, 1.3, 1.3],
        font_size=8.0,
    )
    h2(doc, "7.3 分期实施图表")
    add_table(
        doc,
        ["阶段", "目标", "重点工程", "运营导入", "验收重点"],
        [
            ("近期 0-6个月", "补齐安全底线", "扶手、防滑、照明、座椅、导视、坡道与消防口边界", "护理陪同路线、跌倒风险复盘、老人试走", "无障碍通行、夜间安全、休息点连续性"),
            ("中期 6-18个月", "激活核心节点", "认知体感墙、园艺疗法、地滚球、声景与芳香节点", "园艺课程、康复小组、家属探访活动", "参与人数、停留时长、设施完好率"),
            ("长期 18个月以上", "形成机制闭环", "可替换展陈、植物更新、社区共享接口、活动储物与维护系统", "志愿者、学校、社区、医养结合记录", "年度运营日历、满意度、维护成本"),
        ],
        [0.95, 1.15, 2.0, 1.55, 1.35],
        font_size=8.0,
    )

    h1(doc, "8. 总结")
    para(
        doc,
        "深圳市养老护理院景观改造的关键，不在于增加装饰性绿化，而在于把老人真实、重复、脆弱但有活力的日常行为转化为空间系统。院内通过“五园一廊一径”形成安全可控的日常生活网络，街心公园通过南北分区形成认知友好、运动干预和社区共享界面。"
    )
    para(
        doc,
        "项目实施应坚持三条原则：第一，安全先行，任何疗愈和社交都建立在无障碍、照明、防滑、扶手和可视化看护之上；第二，参与优先，景观不只被观看，还应被触摸、种植、聆听、回忆和共同使用；第三，运营前置，空间节点必须对应护理、康复、社工、家属和社区志愿活动，否则改造难以持续产生健康促进效益。"
    )
    para(
        doc,
        "最终目标是实现三重转变：从“养老空间”转向“生活空间”，从“被动照护”转向“积极参与”，从“封闭机构”转向“健康社区”。"
    )

    h1(doc, "参考文献与资料来源")
    refs = [
        (
            "[R1] 深圳市南山区人民政府：深圳市养老护理院机构信息页。",
            "https://www.szns.gov.cn/bsfw/nsqyljgmc/szsylhly/content/post_9447864.html",
        ),
        (
            "[R2] 深圳市人民政府：深圳市养老服务设施布局专项规划（2025-2035年）相关公开信息。",
            "https://www.sz.gov.cn/cn/xxgk/zfxxgj/ghjh/csgh/zxgh/content/post_12470457.html",
        ),
        (
            "[R3] World Health Organization. Active Ageing: A Policy Framework. 2002.",
            "https://extranet.who.int/agefriendlyworld/wp-content/uploads/2014/06/WHO-Active-Ageing-Framework.pdf",
        ),
        (
            "[R4] World Health Organization. Global Age-friendly Cities: A Guide. 2007.",
            "https://iris.who.int/handle/10665/43755",
        ),
        (
            "[R5] Whole Building Design Guide. Architectural Programming.",
            "https://www.wbdg.org/design-disciplines/architectural-programming",
        ),
        (
            "[R6] Ulrich, R. S. View through a window may influence recovery from surgery. Science, 1984.",
            "https://doi.org/10.1126/science.6143402",
        ),
        (
            "[R7] Marcus, C. C., & Sachs, N. A. Therapeutic Landscapes: An Evidence-Based Approach to Designing Healing Gardens and Restorative Outdoor Spaces. Wiley, 2014.",
            "https://www.wiley.com/en-us/Therapeutic+Landscapes%3A+An+Evidence+Based+Approach+to+Designing+Healing+Gardens+and+Restorative+Outdoor+Spaces-p-9781118421109",
        ),
        (
            "[R8] 住房和城乡建设部：建筑与市政工程无障碍通用规范 GB 55019-2021。",
            "https://www.mohurd.gov.cn/",
        ),
        (
            "[R9] 住房和城乡建设部：老年人照料设施建筑设计标准 JGJ 450-2018。",
            "https://www.mohurd.gov.cn/",
        ),
        (
            "[R10] 住房和城乡建设部：城市居住区规划设计标准 GB 50180-2018。",
            "https://www.mohurd.gov.cn/gongkai/zhengce/zhengcefilelib/201811/20181130_238590.html",
        ),
        (
            "[R11] 源PPT：建筑策划汇报.pptx，2026-05-25，深圳市养老护理院景观改造建筑策划报告。",
            SOURCE_PPT,
        ),
    ]
    for text, url in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text + " ")
        set_run_font(r, size=9, color="1F2937")
        if url.startswith("http"):
            add_hyperlink(p, url, url)
        else:
            r2 = p.add_run(url)
            set_run_font(r2, size=9, color="6B7280")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = make_doc()
    print(path)
