from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "深圳市养老护理院景观改造建筑策划报告_修订版.docx"
SOURCE_PPT = r"C:\Users\19727\Desktop\建筑策划\建筑策划汇报.pptx"


def font(run, size=10.5, bold=False, color="1F2937", name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=110, bottom=80, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def cell_text(cell, text, size=8.5, bold=False, color="1F2937", align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text))
    font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def table(doc, headers, rows, widths, header_fill="E8EEF5", size=8.3):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, width in enumerate(widths):
        t.rows[0].cells[i].width = Inches(width)
        shade(t.rows[0].cells[i], header_fill)
        cell_text(t.rows[0].cells[i], headers[i], size=size, bold=True, color="0B2545", align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        row = t.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].width = Inches(widths[i])
            cell_text(row.cells[i], value, size=size)
    doc.add_paragraph()
    return t


def p(doc, text, lead=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(5)
    para.paragraph_format.line_spacing = 1.12
    if lead and text.startswith(lead):
        r = para.add_run(lead)
        font(r, bold=True, color="0B2545")
        r = para.add_run(text[len(lead):])
        font(r)
    else:
        r = para.add_run(text)
        font(r)
    return para


def h1(doc, text):
    para = doc.add_heading(text, level=1)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(7)
    for r in para.runs:
        font(r, size=16, bold=True, color="2E74B5")


def h2(doc, text):
    para = doc.add_heading(text, level=2)
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(5)
    for r in para.runs:
        font(r, size=12.5, bold=True, color="1F4D78")


def bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(3)
        r = para.add_run(item)
        font(r, size=10)


def hyperlink(paragraph, text, url):
    rid = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    txt = OxmlElement("w:t")
    txt.text = text
    run.append(txt)
    link.append(run)
    paragraph._p.append(link)


def add_refs(doc):
    h1(doc, "参考文献与资料来源")
    refs = [
        ("深圳市南山区人民政府：深圳市养老护理院机构信息页。", "https://www.szns.gov.cn/bsfw/nsqyljgmc/szsylhly/content/post_9447864.html"),
        ("深圳市人民政府：深圳市养老服务设施布局专项规划（2025-2035年）公开信息。", "https://www.sz.gov.cn/cn/xxgk/zfxxgj/ghjh/csgh/zxgh/content/post_12470457.html"),
        ("World Health Organization. Active Ageing: A Policy Framework. 2002.", "https://extranet.who.int/agefriendlyworld/wp-content/uploads/2014/06/WHO-Active-Ageing-Framework.pdf"),
        ("World Health Organization. Global Age-friendly Cities: A Guide. 2007.", "https://iris.who.int/handle/10665/43755"),
        ("Whole Building Design Guide. Architectural Programming.", "https://www.wbdg.org/design-disciplines/architectural-programming"),
        ("Ulrich, R. S. View through a window may influence recovery from surgery. Science, 1984.", "https://doi.org/10.1126/science.6143402"),
        ("Marcus, C. C., & Sachs, N. A. Therapeutic Landscapes. Wiley, 2014.", "https://www.wiley.com/en-us/Therapeutic+Landscapes%3A+An+Evidence+Based+Approach+to+Designing+Healing+Gardens+and+Restorative+Outdoor+Spaces-p-9781118421109"),
        ("住房和城乡建设部：建筑与市政工程无障碍通用规范 GB 55019-2021。", "https://www.mohurd.gov.cn/"),
        ("住房和城乡建设部：老年人照料设施建筑设计标准 JGJ 450-2018。", "https://www.mohurd.gov.cn/"),
        ("住房和城乡建设部：城市居住区规划设计标准 GB 50180-2018。", "https://www.mohurd.gov.cn/gongkai/zhengce/zhengcefilelib/201811/20181130_238590.html"),
        (f"源PPT：建筑策划汇报.pptx，2026-05-25。{SOURCE_PPT}", ""),
    ]
    for idx, (name, url) in enumerate(refs, 1):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(3)
        r = para.add_run(f"[R{idx}] {name} ")
        font(r, size=8.8)
        if url:
            hyperlink(para, url, url)


def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    for name in ["Heading 1", "Heading 2", "List Bullet", "List Number"]:
        styles[name].font.name = "Microsoft YaHei"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("深圳市养老护理院景观改造建筑策划报告（修订版）")
    font(r, size=8, color="6B7280")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("深圳市养老护理院景观改造建筑策划报告")
    font(r, size=22, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("基于积极老龄化理念的疗愈型适老化景观更新研究")
    font(r, size=13, color="2E74B5")

    table(
        doc,
        ["项目", "内容"],
        [
            ("报告性质", "建筑策划/景观更新策划报告（按工作会议、HECTTEAS矩阵与任务书要求修订）"),
            ("研究对象", "深圳市养老护理院院内庭院、街心公园、室外慢行系统与集中活动区域"),
            ("依据资料", f"源PPT：{SOURCE_PPT}；用户补充参考图；公开政策、标准与循证研究文献"),
            ("面积口径", "本阶段未取得CAD测绘面积，任务书以总更新面积 A 的比例和折算公式表达，施工图阶段复核为平方米。"),
        ],
        [1.35, 8.0],
        size=9,
    )
    p(doc, "核心判断：本项目不是单点景观美化，而是以安全底线、疗愈促进、社交参与、认知友好为主线，把院内庭院与街心公园组织为连续、可停留、可参与、可运营的户外康养生活系统。", "核心判断：")

    h1(doc, "1. 选址")
    p(doc, "项目位于深圳市南山区桃源街道龙珠七路与龙苑路交界处的深圳市养老护理院及其相邻街心公园界面。根据公开项目资料，深圳市养老护理院用地面积约1万平方米，建筑面积约4万平方米，床位约800张，具备长期照护、护理康复和医养结合的机构基础。")
    bullets(doc, [
        "选址性质：既有养老机构存量更新，重点在于释放院内外户外空间的康养价值。",
        "空间边界：院内庭院承接日常护理、低强度康复和安静休憩；街心公园承接外出活动、社区共享和认知友好宣教。",
        "策划重点：以慢行可达、护理可视、日照停留、五感疗愈和持续运营为选址评价核心。",
    ])

    h1(doc, "2. 项目背景研究")
    p(doc, "我国养老机构正在从“床位供给”和“被动照护”转向“健康促进、长期参与和尊严生活”。积极老龄化框架强调健康、参与和保障；年龄友好城市框架强调户外空间、交通、社会参与、尊重与包容。对养老护理院而言，户外空间已成为康复训练、情绪调节、认知刺激、家属探访和公共生活的基础设施。")
    table(
        doc,
        ["研究维度", "背景判断", "对本项目的启示"],
        [
            ("政策与趋势", "养老设施从供给型向健康促进型转变。", "景观更新应进入养老机构服务体系，而不是作为绿化附属。"),
            ("场地现状", "院内庭院、街心公园、慢行系统和集中活动区已有高频使用基础。", "以连续网络整合碎片空间，优先处理安全和可达。"),
            ("使用行为", "散步、晒太阳、交谈、锻炼、探访是老人日常高频行为。", "把行为转译为路径、节点、家具、导视和运营活动。"),
            ("设计依据", "疗愈景观、无障碍设计、养老设施标准和年龄友好原则共同约束。", "建立安全性、疗愈性、社交性、认知友好四位一体目标。"),
        ],
        [1.25, 3.8, 4.3],
        size=8.8,
    )

    h1(doc, "3. 收集与分析信息")
    p(doc, "信息收集分为现场与源PPT证据、老人行为需求、运维与会议反馈、政策标准与循证研究四类。由于未获得CAD红线、竖向标高、日照模拟和精确面积，以下分析以策划阶段的问题识别和指标控制为主。")
    table(
        doc,
        ["信息类型", "已获得信息", "分析结论", "后续需补充"],
        [
            ("场地现状", "院内庭院、街心公园、慢行系统、集中活动场地照片与问题归纳。", "存在安全、体验、认知和运营四类缺口。", "CAD、竖向标高、坡度、消防口边界、铺装防滑参数。"),
            ("行为观察", "早餐后、午饭后、傍晚散步；中庭与南向晒太阳；庭院交谈和探访。", "路径、日照、社交、康复和认知刺激是空间组织主线。", "分时段人流、停留热力、跌倒风险和护理陪同记录。"),
            ("使用人群", "自理、介助、介护老人，家属，护理人员，康复师，社工，志愿者。", "空间需分级开放，兼顾自主活动与看护安全。", "老人行动能力分级、认知症比例、护理视线盲区。"),
            ("运营条件", "现状已有活动需求，但固定课程、场景接口和维护机制不足。", "更新策略必须和护理、康复、社工、社区联动。", "年度活动日历、维护预算、志愿者机制。"),
        ],
        [1.1, 3.0, 2.8, 2.4],
        size=8.4,
    )

    h1(doc, "4. 工作会议")
    p(doc, "本章按照用户参考图1和图2重构：先由访谈/会议记录反推价值议题，再形成“价值-目标-事实信息-需求-策划概念”成果矩阵。该矩阵作为后续概念策划、HECTTEAS排序和任务书的输入。")
    h2(doc, "4.1 工作会议记录反推")
    table(
        doc,
        ["价值议题", "目标类发言（反推）", "事实信息类发言（反推）", "需求类发言（反推）", "会议归纳"],
        [
            ("环境", "院方希望“增加一点聚集地，使老人愿意交流”；老人代表提出“晒太阳的地方要舒服”。", "植物生长杂乱、层次不足；导视不清；中庭和南向空间自发停留多；街心公园高差阻隔。", "需要绿荫、日照、休息、芳香植物、清晰导视和低维护植物。", "环境价值应从绿化美化提升为疗愈、停留与识别系统。"),
            ("功能", "护理部提出“空间要服务康复和日常照护”；家属希望探访时能有可坐、可看的地方。", "散步、交谈、锻炼、探访高频发生；现状健身设施基础，集中活动区座椅与日照不足。", "需要低强度康复、地滚球、园艺课程、家属探访、展示活动和护理陪同点。", "功能应复合组织，避免单一通行或单一观赏。"),
            ("交通", "康复师强调老人要“知道往哪走、走累了能停”。", "扶手不连续；台阶和高差不利于轮椅/介助老人；消防口打断街心公园线性空间。", "需要连续慢行、坡道、扶手、50米左右休息节点、夜间照明和人车/内外分流。", "交通系统是康养活动骨架，先解决安全连续，再叠加疗愈节点。"),
            ("经济", "院方与运维希望“施工影响小、维护压力可控”。", "既有机构运营不中断；护理院施工扰动敏感；后期维护人手有限。", "需要分期实施、低成本微更新、模块化设施、低维护种植和可替换展陈。", "经济价值强调短周期、可迭代和可维护，而非一次性重投入。"),
            ("文化", "社工提出“老人需要被看见，也需要和社区有联系”。", "机构封闭感较强；记忆展示、认知症友好宣教和社区参与不足。", "需要老物件、老照片、节庆展示、志愿活动、认知友好宣教和社区共享规则。", "文化价值应转译为记忆疗愈和健康社区界面。"),
            ("形式", "设计团队提出空间要“温暖、可识别、有主题”。", "现状流线固定，趣味性不足；节点记忆弱；部分材料触感和家具适老性不足。", "需要主题花园、色彩导视、适老座椅、触觉材料、声景和清晰边界。", "形式不是装饰，而是帮助老人识别、停留和产生参与意愿的媒介。"),
        ],
        [0.8, 2.25, 2.4, 2.25, 1.8],
        header_fill="F2F4F7",
        size=7.5,
    )
    h2(doc, "4.2 工作会议成果矩阵")
    table(
        doc,
        ["价值", "目标", "事实信息", "需求", "策划概念"],
        [
            ("环境", "营造疗愈景观；提升日照与树荫停留品质；强化空间识别。", "植被维护粗放、层次不足；导视不清；南向和中庭停留频繁；街心公园活力不足。", "连续绿荫、冬季日照、芳香/声景、休息座椅、低维护植物。", "五感疗愈花园、翠阴廊、凝神之园、芳香与声景节点。"),
            ("功能", "满足生活照护、康复运动、家属探访和社区共享。", "散步、晒太阳、交谈、锻炼、探访高频；互动设施不足；集中活动场地不够友好。", "地滚球、园艺课程、探访座椅、活动展示、护理陪同与健康观察。", "康竞园、丰收径、共聚之园、八方园、家属探访节点。"),
            ("交通", "形成安全连续、可识别、可休息的慢行系统。", "扶手不连续；台阶/高差阻隔；消防口打断；路径识别弱。", "无障碍坡道、连续扶手、防滑铺装、50米左右休息点、夜间照明。", "探索康养径、内外分流、图形导视、节点打卡系统。"),
            ("经济", "低成本、短周期、可分期、可维护。", "存量场地更新；施工影响敏感；维护力量有限；设施需长期完好。", "针灸式微更新、模块化设施、可替换展陈、低维护种植。", "近期安全底线包、中期节点激活包、长期运营维护包。"),
            ("文化", "降低机构封闭感；增强记忆疗愈与社区连接。", "记忆展陈不足；认知友好宣传弱；家属与志愿者参与场景不足。", "老物件/老照片、认知症友好宣教、节庆展示、社区共建。", "忆录园、认知体感墙、社区共享界面、志愿活动花园。"),
            ("形式", "形成适老尺度、温暖美观、易识别的空间形象。", "流线固定、趣味性不足、节点记忆弱、家具与材料适老性不足。", "主题化节点、色彩分区、触觉材料、半围合座椅、明确边界。", "五园一廊一径；街心公园南北两区；可触摸、可参与的景观节点。"),
        ],
        [0.75, 2.0, 2.35, 2.15, 2.1],
        header_fill="DDEBF7",
        size=7.7,
    )

    h1(doc, "5. 概念性策划")
    p(doc, "总体概念为“探索康养径串联的疗愈型适老化户外生活网络”。院内以“五园一廊一径”承接日常生活、康复和记忆疗愈；街心公园以南向认知互动与宣教区、北向运动干预与社交活动区承接外出活动和社区共享。")
    table(
        doc,
        ["空间层级", "组成", "核心功能", "控制要点"],
        [
            ("一径", "探索康养径", "连续步行、康复训练、节点串联、展示参观。", "闭环或准闭环；按无障碍标准复核坡度、净宽、扶手和防滑。"),
            ("一廊", "翠阴廊", "阴凉通行、植物疗愈、慢行停留。", "攀藤植物、遮阴、可识别导视和适老座椅。"),
            ("五园", "八方园、三松园、隐逸园、康竞园、忆录园，结合丰收径园艺系统。", "入口聚集、活动展示、安静休憩、低强度运动、记忆疗愈、园艺课程。", "每个节点有明确主题、借力座椅、护理可视范围和运营活动。"),
            ("街心公园", "南向认知互动与宣教区；北向运动干预与社交活动区。", "认知体感、科普展示、地滚球、平衡训练、社区共聚。", "处理高差、消防口打断和可管理开放边界。"),
        ],
        [0.9, 2.1, 3.0, 3.0],
        size=8.1,
    )
    bullets(doc, [
        "路径缝合：把碎片空间串联为连续康养径，形成可达、可停、可看护的户外活动骨架。",
        "针灸式微更新：优先更新高频节点，以模块化设施和低维护种植降低施工与运营压力。",
        "多感官疗愈：将视觉、触觉、听觉、嗅觉、味觉和运动觉落到植物、铺装、声景、园艺和康复活动中。",
        "场景复合：把散步、社交、宣教、园艺、康复运动和家属探访叠合在日常路径上。",
    ])

    h1(doc, "HECTTEAS 矩阵分析")
    p(doc, "参考图3重新绘制 HECTTEAS 矩阵。矩阵以“价值排序”为起点，横向展开目标、事实、需求、问题和概念，用于把会议结论转译为建筑策划任务。")
    table(
        doc,
        ["价值 VALUES", "目标 GOALS", "事实 FACTS", "需求 NEEDS", "问题 PROBLEMS", "概念 IDEAS"],
        [
            ("人文因素\nHUMAN", "让老人从被动照护转向积极参与，兼顾自理、介助、介护差异。", "散步、日照、交谈、锻炼、探访高频发生；部分老人方向感和行动能力下降。", "可达路径、可停座椅、护理可视、家属探访、低压力社交。", "空间支持不足会降低外出意愿，并加重护理陪同压力。", "探索康养径、半围合交流节点、家属探访花园、护理观察点。"),
            ("环境因素\nENVIRONMENTAL", "构建自然疗愈、热舒适和低维护的户外环境。", "植物层次和维护不足；树荫与日照分布不均；街心公园停留活力弱。", "遮阴、冬季日照、芳香植物、声景、可食植物和雨水友好铺装。", "只做绿化会停留在观看层面，无法支撑持续疗愈。", "五感花园、翠阴廊、芳香/声景节点、低维护植物系统。"),
            ("文化因素\nCULTURAL", "增强记忆疗愈、机构认同和社区共享。", "认知友好宣教不足；老人生活记忆和社区连接未被空间化。", "老照片、老物件、节庆展示、志愿活动、认知症友好宣传。", "机构封闭感强，家属与社区参与缺少稳定场景。", "忆录园、认知体感墙、社区共享界面、多代互动活动。"),
            ("技术因素\nTECHNICAL", "落实无障碍、导视、照明、求助和运维技术控制。", "扶手不连续；高差与台阶阻隔；导视识别弱；消防口打断线性空间。", "坡道、扶手、防滑、照明、紧急求助、图形导视、消防边界。", "技术系统不连续会直接形成跌倒、迷失和运营风险。", "连续安全设施包、色彩分区导视、夜间照明和求助系统。"),
            ("时间因素\nTEMPORAL", "分期实施，减少施工扰动并让运营逐步导入。", "既有养老机构运营不中断；改造涉及老人日常路径。", "近期补安全，中期激活节点，长期形成活动与维护闭环。", "一次性改造容易影响照护秩序，也难以根据反馈迭代。", "0-6个月安全底线包、6-18个月节点激活包、18个月以上运营固化。"),
            ("经济因素\nECONOMIC", "控制成本、维护压力和后期更换成本。", "存量更新条件下造价边界明确；维护人力有限。", "模块化设施、可替换展陈、低维护种植、志愿与社区资源。", "高维护景观会增加院方长期负担。", "针灸式微更新、可替换展陈、设施标准化采购与维护清单。"),
            ("美学因素\nAESTHETIC", "形成温暖、清晰、可识别、有尊严的养老空间形象。", "现状趣味性和节点记忆不足，空间形象偏功能化。", "主题花园、季相色彩、触觉材料、适老家具、界面展示。", "美学若只停留在装饰，会与老人真实行为脱节。", "五园一廊一径主题系统、南北分区形象、可参与景观。"),
            ("安全因素\nSAFETY", "把跌倒预防、迷失预防、夜间安全和消防边界作为底线。", "高差、台阶、扶手断点、照明和消防口打断是主要风险。", "连续扶手、防滑铺装、照明、休息点、紧急求助、护理可视。", "安全底线不稳，疗愈、社交和开放共享都无法成立。", "安全闭环路径、50米左右休息节点、可视看护点和风险巡检表。"),
        ],
        [1.05, 1.85, 2.0, 1.85, 1.85, 1.95],
        header_fill="D9EAD3",
        size=7.1,
    )

    h1(doc, "6. 空间分配表")
    p(doc, "空间分配采用“比例控制+面积公式”。A 为施工图阶段经 CAD/测绘复核后的总更新面积，表中面积为 A 的折算值。")
    table(
        doc,
        ["空间类型", "比例", "面积", "主要位置", "设施与内容", "优先级"],
        [
            ("连续康养径与无障碍慢行", "28%", "0.28A", "院内路径、街心公园衔接段", "环线、扶手、防滑铺装、坡道、导视、50米左右休息点", "高"),
            ("日照与安静休憩", "14%", "0.14A", "中庭、南向停留点、隐逸园", "半围合座椅、可移动家具、遮阴与冬季日照平衡", "高"),
            ("多感官疗愈花园", "18%", "0.18A", "凝神之园、共聚之园、翠阴廊、芳香节点", "季相色彩、触觉材料、声景、芳香植物、冥想停留", "高"),
            ("社交探访与展示活动", "12%", "0.12A", "八方园、三松园、入口与街心界面", "家属探访座椅、活动展示、节庆布置、社区宣教", "中高"),
            ("康复运动与地滚球", "10%", "0.10A", "康竞园、街心公园北向区域", "环形跑道、地滚球、平衡训练、护理陪同点", "中高"),
            ("园艺疗法与共享种植", "8%", "0.08A", "丰收径、共享种植节点", "抬高种植床、采摘工具、可食植物、园艺课程", "中"),
            ("导视照明求助与管理设施", "6%", "0.06A", "全场关键节点", "色彩导视、图形标识、照明、紧急求助、维护储物", "高"),
            ("维护缓冲与弹性机动", "4%", "0.04A", "边界、消防口、临时活动区", "消防边界、可替换展陈、临时活动接口", "中"),
        ],
        [1.6, 0.55, 0.65, 1.6, 3.2, 0.55],
        size=7.6,
    )

    h1(doc, "7. 空间图表")
    table(
        doc,
        ["高频行为", "空间需求", "设施落位", "运营活动"],
        [
            ("散步", "连续、可识别、可休息的慢行环线。", "探索康养径、扶手、遮阴、休息节点、里程打卡。", "晨间散步、饭后慢行、康复步行记录。"),
            ("晒太阳", "南向、半开放、可移动座椅和冬季日照。", "中庭、南向停留点、可移动椅、暖色铺装。", "午间日照陪伴、家属探访、节气活动。"),
            ("交谈", "小尺度围合、可面对面、护理人员可观察。", "八方园、三松园、共聚之园、家属座椅。", "茶话会、生日会、志愿陪伴。"),
            ("锻炼", "低强度、低冲击、可陪同与可中止。", "康竞园、地滚球、平衡训练、环形跑道。", "康复训练、地滚球小组、健康观察记录。"),
            ("认知刺激", "低复杂度、可触摸、可回忆、可重复。", "认知体感墙、老物件展示、色彩导视、芳香植物。", "记忆叙事、认知症友好宣教、家属共创展陈。"),
        ],
        [1.0, 2.4, 3.0, 2.9],
        size=8.0,
    )
    table(
        doc,
        ["入口/展示", "日常康养路径", "安静疗愈", "参与活动", "外向共享"],
        [
            ("八方园\n门面聚集、树下座椅、活动展示", "探索康养径\n串联院内节点与街心公园", "隐逸园\n绿篱隔离、安静休憩", "丰收径\n种植箱、园艺课程、采摘互动", "街心公园南向\n认知互动与宣教"),
            ("三松园\n灵活家具、方院活动、记忆节点", "翠阴廊\n拱形绿廊、阴凉通行", "凝神之园\n五感安抚、冥想停留", "康竞园\n环形跑道、地滚球、平衡训练", "街心公园北向\n运动干预与社区共聚"),
            ("忆录园\n老物件展示、记忆疗愈", "导视系统\n图形识别、色彩引导", "芳香/声景节点\n嗅觉与听觉疗愈", "共聚之园\n家属探访、志愿活动", "社区界面\n可管理开放、养老形象展示"),
        ],
        [1.86, 1.86, 1.86, 1.86, 1.86],
        header_fill="FCE4D6",
        size=7.8,
    )

    h1(doc, "8. 总结")
    p(doc, "深圳市养老护理院景观改造的关键，不在于增加装饰性绿化，而在于把老人真实、重复、脆弱但有活力的日常行为转化为空间系统。院内通过“五园一廊一径”形成安全可控的日常生活网络，街心公园通过南北分区形成认知友好、运动干预和社区共享界面。")
    p(doc, "项目实施应坚持三条原则：安全先行、参与优先、运营前置。最终目标是实现从“养老空间”到“生活空间”、从“被动照护”到“积极参与”、从“封闭机构”到“健康社区”的转变。")

    h1(doc, "9. 建筑策划任务书")
    p(doc, "本任务书作为后续方案设计、初步设计和施工图深化的控制依据。面积以 A 表示总更新面积，待测绘/CAD复核后换算为平方米。")
    table(
        doc,
        ["功能分区", "内容", "面积", "要求"],
        [
            ("A区：院内探索康养径", "连续慢行环线、扶手、防滑铺装、休息节点、导视、照明、护理观察点。", "0.28A", "按GB 55019等无障碍要求复核；休息点约50米设置；路径避免无效折返和护理盲区。"),
            ("B区：院内五园节点", "八方园、三松园、隐逸园、康竞园、忆录园及节点家具。", "0.24A", "每园形成明确主题、停留方式和运营活动；座椅需便于借力、轮椅停靠和护理陪同。"),
            ("C区：翠阴廊与多感官疗愈系统", "绿廊、芳香植物、触觉材料、声景装置、季相植物、冥想停留。", "0.18A", "植物低维护、无明显毒刺风险；声景控制音量和突发噪声；兼顾遮阴与通风。"),
            ("D区：丰收径与园艺疗法", "抬高种植床、可食植物、采摘工具、园艺课程接口、储物设施。", "0.08A", "种植床满足站姿/坐姿操作；工具安全可控；课程与护理、社工排期联动。"),
            ("E区：街心公园南向认知互动与宣教", "认知体感墙、老照片/老物件展陈、科普互动、低复杂度导视。", "0.10A", "内容避免信息过载；图形、色彩和文字高度适老；兼顾对外展示与院方管理。"),
            ("F区：街心公园北向运动干预与社交", "地滚球、平衡训练、合奏之园、社区共聚座椅、活动边界。", "0.12A", "运动设施低冲击、可陪同、可中止；处理消防口和高差；避免与城市通行冲突。"),
            ("G区：导视照明求助与运维系统", "色彩导视、图形标识、夜间照明、紧急求助、维护储物、可替换展陈。", "0.06A", "建立点位清单和巡检制度；照明、求助、消防、物业维护同步审查。"),
            ("H区：弹性与缓冲空间", "消防口边界、设备缓冲、临时活动接口、施工转换区。", "0.04A", "不得侵占消防和必要通行；支持节庆活动、志愿活动和后期设施迭代。"),
        ],
        [1.75, 3.1, 0.7, 3.75],
        header_fill="E2F0D9",
        size=7.7,
    )

    add_refs(doc)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(make_doc())
